import argparse
from pathlib import Path

from docx import Document

from src import db

OUTPUT_DIR = Path(__file__).resolve().parent.parent / "output"


def export_round(run_id: int, round_: int, output_path: Path):
    conn = db.get_connection()

    run_row = db.get_run(conn, run_id)
    if run_row is None:
        raise SystemExit(f"No run with id {run_id}")

    topics = db.get_topics(conn)
    models = db.get_models(conn)

    doc = Document()
    title = "Metanoia Council — Statements"
    doc.add_heading(title, level=0)

    label = run_row["label"] or f"Run {run_id}"
    doc.add_paragraph(f"{label}  ·  round {round_}  ·  started {run_row['started_at']}")

    for topic in topics:
        doc.add_heading(topic["name"], level=1)
        statements = db.get_statements_for_round(conn, run_id, topic["id"], round_)

        if not statements:
            doc.add_paragraph("(no statements recorded for this topic/round)")
            continue

        for model in models:
            stmt = statements.get(model["id"])
            if stmt is None:
                continue

            heading = doc.add_paragraph()
            run = heading.add_run(f"{model['name']} ({model['lab']})")
            run.bold = True

            doc.add_paragraph(stmt["text"])

            if stmt["rationale"]:
                rationale_p = doc.add_paragraph()
                rationale_run = rationale_p.add_run(f"Revision rationale: {stmt['rationale']}")
                rationale_run.italic = True

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path


def main():
    parser = argparse.ArgumentParser(description="Export council statements to a .docx file")
    parser.add_argument("--run", type=int, required=True, help="Run id to export")
    parser.add_argument("--round", type=int, default=1, help="Round to export (default 1)")
    parser.add_argument("--output", type=Path, help="Output .docx path")
    args = parser.parse_args()

    output_path = args.output or (OUTPUT_DIR / f"run{args.run}_round{args.round}_statements.docx")
    saved = export_round(args.run, args.round, output_path)
    print(f"Saved: {saved}")


if __name__ == "__main__":
    main()
